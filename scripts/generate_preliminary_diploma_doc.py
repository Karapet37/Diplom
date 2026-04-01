from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = REPO_ROOT / 'docs' / 'karapet_karapetyan_preliminary_diploma_hy.md'
OUTPUT_PATH = REPO_ROOT / 'ԿարապետԿարապետյան․docx'


def _set_base_fonts(document: Document) -> None:
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for style_name in ('Title', 'Heading 1', 'Heading 2', 'Heading 3'):
        style = document.styles[style_name]
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    document.styles['Title'].font.size = Pt(18)
    document.styles['Heading 1'].font.size = Pt(15)
    document.styles['Heading 2'].font.size = Pt(13)
    document.styles['Heading 3'].font.size = Pt(12)


def _append_paragraph(document: Document, text: str, *, style: str | None = None, center: bool = False) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text.strip())
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(12 if style not in {'Title', 'Heading 1', 'Heading 2', 'Heading 3'} else run.font.size.pt if run.font.size else 12)
    paragraph.paragraph_format.space_after = Pt(6)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _render_markdown(document: Document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    first_title_rendered = False
    previous_blank = True

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            previous_blank = True
            continue

        if stripped.startswith('# '):
            text = stripped[2:].strip()
            if not first_title_rendered:
                paragraph = document.add_paragraph(style='Title')
                run = paragraph.add_run(text)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(10)
                first_title_rendered = True
            else:
                _append_paragraph(document, text, style='Heading 1')
            previous_blank = False
            continue

        if stripped.startswith('## '):
            _append_paragraph(document, stripped[3:].strip(), style='Heading 1')
            previous_blank = False
            continue

        if stripped.startswith('### '):
            _append_paragraph(document, stripped[4:].strip(), style='Heading 2')
            previous_blank = False
            continue

        if stripped.startswith('- '):
            _append_paragraph(document, stripped[2:].strip(), style='List Bullet')
            previous_blank = False
            continue

        if re.match(r'^\d+\.\s+', stripped):
            payload = re.sub(r'^\d+\.\s+', '', stripped)
            _append_paragraph(document, payload, style='List Number')
            previous_blank = False
            continue

        if stripped.startswith('Հեղինակ՝') or stripped == 'Նախնական տարբերակ':
            _append_paragraph(document, stripped, center=True)
            previous_blank = False
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(18)
        paragraph.paragraph_format.space_after = Pt(6)
        if previous_blank:
            paragraph.paragraph_format.space_before = Pt(2)
        run = paragraph.add_run(stripped)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)
        previous_blank = False


def generate() -> Path:
    if not SOURCE_PATH.exists():
        raise FileNotFoundError(f'Source markdown not found: {SOURCE_PATH}')

    markdown_text = SOURCE_PATH.read_text(encoding='utf-8')
    document = Document()
    _set_base_fonts(document)
    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    _render_markdown(document, markdown_text)
    document.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


if __name__ == '__main__':
    path = generate()
    print(path)
